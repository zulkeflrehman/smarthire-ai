"""
parser.py — Extract plain text from PDF and DOCX resume files.

Uses:
  pdfplumber  → PDF  (handles multi-column layouts better than PyPDF2)
  python-docx → DOCX
"""
import pdfplumber
import docx
from pathlib import Path


def _extract_pdf(path: str) -> str:
    """Extract text from a PDF, page by page."""
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n".join(pages)


def _extract_docx(path: str) -> str:
    """Extract text from a DOCX (paragraphs + tables)."""
    doc = docx.Document(path)
    lines = []

    # Body paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text)

    # Tables (often used in resume layouts)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text)

    return "\n".join(lines)


def extract_text(file_path: str) -> str:
    """
    Route file to the correct parser based on extension.

    Returns:
        str: extracted plain text (UTF-8, newline-delimited)

    Raises:
        ValueError: if file type is unsupported or parse fails
    """
    ext = Path(file_path).suffix.lower()
    try:
        if ext == '.pdf':
            text = _extract_pdf(file_path)
        elif ext in ('.docx', '.doc'):
            text = _extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: '{ext}'. Upload PDF or DOCX only.")
    except Exception as exc:
        raise ValueError(f"Could not parse '{Path(file_path).name}': {exc}") from exc

    if not text.strip():
        raise ValueError(f"No text found in '{Path(file_path).name}'. The file may be image-only or empty.")

    return text.strip()
